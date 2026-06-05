# interfaz_programa.py — v53.0 (EXCEL ALINEACION MIXTA CORREGIDA + PDF PERFECTO)
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog, Menu
import pandas as pd
import sqlite3, os, re, sys, traceback, json, difflib
from pathlib import Path
from datetime import date, datetime

# --- Excel ---
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.worksheet.page import PageMargins
from openpyxl.worksheet.pagebreak import Break
try:
    from openpyxl.drawing.image import Image as OXLImage
    HAS_OXL_IMAGE = True
except ImportError:
    HAS_OXL_IMAGE = False

# --- Word ---
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# =============================================================================
#  SECCIÓN 1: CONFIGURACIÓN
# =============================================================================

def app_dir() -> Path:
    if getattr(sys, "frozen", False): return Path(sys.executable).parent
    return Path(__file__).parent

BASE_DIR = app_dir(); DATA_DIR = BASE_DIR / "data"; ASSETS_DIR = BASE_DIR / "assets"
PROGRAMAS_DIR = BASE_DIR / "programas"; RESULTADOS_DIR = BASE_DIR / "resultados"
DATA_DIR.mkdir(exist_ok=True); ASSETS_DIR.mkdir(exist_ok=True)
PROGRAMAS_DIR.mkdir(exist_ok=True); RESULTADOS_DIR.mkdir(exist_ok=True)
DB_PATH = DATA_DIR / "carreras.db"; NOMBRE_BD = str(DB_PATH)
REGISTRO_PATH = DATA_DIR / "cargas_registro.json"

programa_completo = [] 
indice_edicion = None 
DATOS_WORD_CACHED = []

# Variables GUI
entry_fecha = None; entry_nro_reunion = None; entry_nro_carrera = None; entry_horario = None; entry_premio = None
entry_distancia = None; entry_condicion = None; entry_premios_dinero = None; entry_apuesta = None
entry_incremento = None; entry_incremento_2 = None; combo_word = None; combo_dist = None
text_caballos = None; text_kilos = None; text_actuaciones = None; tabla_programa = None; lista_carreras = None
contador_carreras = None; btn_accion = None; estado_db_var = None

# --- COLORES OFICIALES MANDILES ---
MANDILES = {
    "1": ("#d32f2f", "#ffffff"), "2": ("#ffffff", "#000000"), "3": ("#1565c0", "#ffffff"),
    "4": ("#fdd835", "#000000"), "5": ("#2e7d32", "#ffffff"), "6": ("#000000", "#fff200"),
    "7": ("#ef6c00", "#000000"), "8": ("#f48fb1", "#000000"), "9": ("#00bcd4", "#000000"),
    "10": ("#7b1fa2", "#ffffff"), "11": ("#9e9e9e", "#da2128"), "12": ("#71bf44", "#000000"),
    "13": ("#a05b09", "#fff200"), "14": ("#b71c1c", "#ffffff"), "15": ("#f3d19c", "#000000"),
    "default": ("#CCCCCC", "#000000")
}

RECORDS = {
    "700":  '700 metros - Record Dist.: 38" 2/5, Holly 28/02/2026',
    "800":  '800 metros - Record Dist.: 43" 1/5, Enloqueceme 14/12/2025',
    "1000": '1.000 metros - Record Dist.: 58" 2/5, Sarfo 23/07/2020',
    "1100": '1.100 metros - Record Dist.: 1\' 04" 3/5, Sold Out 29/08/2021',
    "1200": '1.200 metros - Record Dist.: 1\' 09" 4/5, Donald Music 28/11/2021',
    "1300": '1.300 metros - Record Dist.: 1\' 16" 2/5, Panatta 19/12/2021 - High Commander 08/10/2023 - Jolly Boy 30/11/2025',
    "1400": '1.400 metros - Record Dist.: 1\' 22" 3/5, Patani 29/10/2017 - Dipinto 28/11/2021',
    "1500": '1.500 metros - Record Dist.: 1\' 28" 4/5, Storm Chuck 24/08/2025',
    "1600": '1.600 metros - Record Dist.: 1\' 37" 2/5, Batman Crest 17/08/1975 - Teenek 11/06/2021 - Latan Craf 22/06/2025',
    "1800": '1.800 metros - Record Dist.: 1\' 51" 4/5, Sir Melody 26/11/2017',
    "2000": '2.000 metros - Record Dist.: 2\' 03" 2/5, Volynov 1978',
    "2200": '2.200 metros - Record Dist.: 2\'17"1/5, Frances Net 24/09/2016',
}

COLORS = {"bg": "#f0f2f5", "primary": "#248689", "accent": "#f16536", "card": "#ffffff", "ink": "#1f2937", "line": "#e5e7eb"}

# =============================================================================
#  SECCIÓN 2: LÓGICA
# =============================================================================

def formatear_cuerpos(valor):
    s = str(valor).strip()
    if any(x in s.upper() for x in ['CZA', 'PZO', 'HCO', 'S.A']): return s
    try:
        # Soporta "2 1/2", "3/4", "1/2" además de decimales como "2.5"
        if '/' in s:
            parts = s.split()
            if len(parts) == 2:
                num, den = parts[1].split('/')
                f = int(parts[0]) + int(num) / int(den)
            else:
                num, den = s.split('/')
                f = int(num) / int(den)
        else:
            f = float(s)
        entero = int(f); dec = f - entero; frac = ""
        if abs(dec - 0.25) < 0.01: frac = "1/4"
        elif abs(dec - 0.50) < 0.01: frac = "1/2"
        elif abs(dec - 0.75) < 0.01: frac = "3/4"
        if entero > 0 and frac: res = f"{entero} {frac}"
        elif entero == 0 and frac: res = frac
        elif entero > 0: res = str(entero)
        else: res = str(s)
        return f"{res} cp"
    except: return s

def _leer_registro():
    if not REGISTRO_PATH.exists():
        return {"programas": {}, "resultados": {}}
    try:
        with open(REGISTRO_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    except: return {"programas": {}, "resultados": {}}

def _guardar_en_registro(tipo, fecha, nombre_archivo):
    reg = _leer_registro()
    reg[tipo][fecha] = nombre_archivo
    with open(REGISTRO_PATH, 'w', encoding='utf-8') as f:
        json.dump(reg, f, ensure_ascii=False, indent=2)

PREFS_PATH = DATA_DIR / "preferencias.json"

def _leer_prefs():
    if not PREFS_PATH.exists(): return {}
    try:
        with open(PREFS_PATH, 'r', encoding='utf-8') as f: return json.load(f)
    except: return {}

def _guardar_prefs(updates: dict):
    prefs = _leer_prefs()
    prefs.update(updates)
    with open(PREFS_PATH, 'w', encoding='utf-8') as f:
        json.dump(prefs, f, ensure_ascii=False, indent=2)

def _leer_estado_db():
    try:
        conn = sqlite3.connect(NOMBRE_BD); c = conn.cursor()
        # Parsear fechas DD-MM-YY correctamente (MAX() de texto falla con este formato)
        rows = c.execute('SELECT DISTINCT snapshot_programa_fecha FROM caballos WHERE snapshot_programa_fecha IS NOT NULL').fetchall()
        res  = c.execute('SELECT MAX(fecha) FROM actuaciones').fetchone()[0]
        conn.close()
        prog_fmt = 'N/D'
        if rows:
            fechas = []
            for (d,) in rows:
                try: fechas.append(datetime.strptime(d, '%d-%m-%y'))
                except: pass
            if fechas:
                prog_fmt = max(fechas).strftime('%d/%m/%y')
        if res:
            try:
                dt = datetime.fromisoformat(str(res).split()[0])
                res_fmt = dt.strftime('%d/%m/%y')
            except: res_fmt = str(res)[:10]
        else: res_fmt = 'N/D'
        return f"Programa: {prog_fmt}  |  Resultados: {res_fmt}"
    except: return "Sin datos en DB"

def _actualizar_estado_db():
    global estado_db_var
    if estado_db_var is not None:
        estado_db_var.set(_leer_estado_db())

def _inicializar_db_si_no_existe():
    conn = sqlite3.connect(NOMBRE_BD); c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS caballos (nombre TEXT PRIMARY KEY, padre_madre TEXT, pelo TEXT, ultima_edad TEXT, ultimo_peso TEXT, ultimo_jockey TEXT, caballeriza TEXT, cuidador TEXT, ultima_actuacion_externa TEXT, texto_actuaciones_externas TEXT, snapshot_programa_fecha DATE)''')
    c.execute('''CREATE TABLE IF NOT EXISTS actuaciones (id INTEGER PRIMARY KEY, fecha DATE, nombre_caballo TEXT, puesto_original INTEGER, puesto_final TEXT, jockey TEXT, cuerpos TEXT, ganador TEXT, segundo TEXT, margen TEXT, tiempo_ganador TEXT, pista TEXT, fue_distanciado BOOLEAN, observacion TEXT)''')
    try: c.execute('ALTER TABLE caballos ADD COLUMN texto_actuaciones_externas TEXT')
    except: pass
    conn.commit(); conn.close()

def conectar_y_cargar_datos():
    _inicializar_db_si_no_existe(); conn = sqlite3.connect(NOMBRE_BD)
    try: df_cab = pd.read_sql_query("SELECT * FROM caballos", conn); df_act = pd.read_sql_query("SELECT * FROM actuaciones", conn)
    except: df_cab = pd.DataFrame(); df_act = pd.DataFrame()
    finally: conn.close()
    if not df_cab.empty: df_cab = df_cab.rename(columns={'nombre':'Caballo', 'ultima_edad':'Edad', 'ultimo_peso':'Peso', 'ultimo_jockey':'Jockey-Descargo', 'padre_madre':'Padre - Madre', 'caballeriza':'Caballeriza', 'cuidador':'Cuidador', 'pelo':'Pelo'})
    if not df_act.empty: 
        df_act = df_act.rename(columns={'nombre_caballo':'Caballo', 'puesto_original':'Puesto Original', 'puesto_final':'Puesto Final', 'jockey':'Jockey', 'cuerpos':'Cuerpos al Ganador', 'ganador':'Ganador', 'segundo':'Segundo', 'margen':'Margen', 'tiempo_ganador':'Tiempo Ganador', 'pista':'Pista', 'fue_distanciado':'Fue Distanciado', 'fecha':'Fecha', 'observacion':'Observacion'})
        df_act['Fecha'] = pd.to_datetime(df_act['Fecha'], errors='coerce')
    return df_cab, df_act

def obtener_datos_caballo(nombre, db_cab, db_act):
    nombre = nombre.strip().upper()
    try: info = db_cab[db_cab['Caballo'] == nombre].iloc[0].to_dict()
    except: info = {'Caballo': nombre}
    
    acts = db_act[db_act['Caballo'] == nombre].sort_values(by='Fecha', ascending=True) 
    
    # --- 1. Calcular las siglas Locales (Tucumán) en orden cronológico ---
    ult_locales = []
    for _, row_act in acts.iterrows():
        ps = str(row_act['Puesto Final']).strip()
        tiempo = str(row_act.get('Tiempo Ganador', ''))
        
        if ps.isdigit():
            puesto = int(ps)
            if puesto >= 10:
                ult_locales.append('0')
            elif puesto == 1:
                # --- MAGIA DEL EXTRAOFICIAL (1e) ---
                es_extraoficial = False
                # Si no tiene el apóstrofe de los minutos (ej. 1'13")
                if "'" not in tiempo:
                    # Buscamos los segundos antes de la comilla doble (ej. 44")
                    m_sec = re.search(r"(\d{2})\"", tiempo)
                    if m_sec:
                        segundos = int(m_sec.group(1))
                        # Si los segundos están entre 36 y 50, es de 700 u 800 mts
                        if 36 <= segundos <= 50:
                            es_extraoficial = True
                
                ult_locales.append('1e' if es_extraoficial else '1')
                # -----------------------------------
            else:
                ult_locales.append(str(puesto))
        else:
            ult_locales.append('-')
            
# --- 2. Traer las siglas Externas (Buenos Aires, etc) ---
    ext_str = str(info.get('ultima_actuacion_externa', '')).strip()
    ult_brutas = [x.strip() for x in ext_str.split('-') if x.strip()] if ext_str and ext_str.lower() != 'nan' else []
    
    # --- 3. Combinar manteniendo orden cronológico ---
    def _es_ext(a):
        return any(c.isalpha() and c.lower() != 'e' for c in a)

    hay_externos = any(_es_ext(a) for a in ult_brutas)

    if not hay_externos:
        # Sin externos: matching hacia adelante para de-duplar pre-BD
        db_ptr = 0
        combined_full = []
        for act in ult_brutas:
            if db_ptr < len(ult_locales) and ult_locales[db_ptr] == act:
                combined_full.append(ult_locales[db_ptr]); db_ptr += 1
            else:
                combined_full.append(act)
        combined_full.extend(ult_locales[db_ptr:])
    else:
        # Con externos: matching desde la DERECHA de ult_locales para encontrar
        # el punto correcto de inserción y preservar el orden cronológico.
        # Los locales de ult_brutas pueden ser recientes (en DB) o pre-BD.
        locals_idx = [(i, a) for i, a in enumerate(ult_brutas) if not _es_ext(a)]
        match_pos = {}  # índice en ult_brutas → índice en ult_locales (o None)
        search_end = len(ult_locales)
        for bi, bact in reversed(locals_idx):
            found = -1
            for k in range(search_end - 1, -1, -1):
                if ult_locales[k] == bact:
                    found = k; break
            match_pos[bi] = found if found >= 0 else None
            if found >= 0:
                search_end = found

        combined_full = []
        db_ptr = 0
        for i, act in enumerate(ult_brutas):
            if _es_ext(act):
                combined_full.append(act)
            else:
                mp = match_pos.get(i)
                if mp is not None and mp >= db_ptr:
                    combined_full.extend(ult_locales[db_ptr:mp + 1])
                    db_ptr = mp + 1
                else:
                    combined_full.append(act)  # pre-BD real
        combined_full.extend(ult_locales[db_ptr:])
    
    # --- 4. Aplicar lógica de prolijidad (Máximo 3 o 4) ---
    # Tomamos las últimas 4 por defecto
    total_ult = combined_full[-4:] if combined_full else []
    
    # Verificamos si en esas 4 hay siglas de afuera (letras que NO sean la 'e' de extraoficial)
    # Por ejemplo, si detecta 'P' o 'LP', tiene_afuera será True. Si solo hay '1e' o números, será False.
    tiene_afuera = any(c.isalpha() and c.lower() != 'e' for act in total_ult for c in act)
    
    # Si detecta que hay actuaciones de afuera en ese grupo, lo recorta a las 3 más nuevas
    if tiene_afuera:
        total_ult = combined_full[-3:] if combined_full else []
        
    cuatro = "-".join(total_ult) if total_ult else "Debuta"
    
    # Guardamos los datos en el diccionario final
    edad = info.get('Edad', '')
    info['E Kg'] = f"{edad} {info.get('Peso','')}".strip()
    info['4 Ult.'] = cuatro
    info['actuaciones'] = acts.tail(2)
    info['texto_act_ext'] = str(info.get('texto_actuaciones_externas', '')).strip()
    # Fechas de todas las actuaciones en DB (para deduplicar con externas)
    try:
        info['todas_act_fechas'] = set(
            acts['Fecha'].dropna().apply(lambda x: x.strftime('%d/%m/%y'))
        )
    except Exception:
        info['todas_act_fechas'] = set()
    
    return info

def cargar_word_entrada():
    if not HAS_DOCX:
        messagebox.showerror("Módulo faltante", "Ejecutar en consola:\npip install python-docx"); return
    f = filedialog.askopenfilename(filetypes=[("Carta de Llamada Word", "*.docx")])
    if not f: return
    try:
        doc = docx.Document(f)
    except Exception as ex:
        messagebox.showerror("Error al leer Word", f"No se pudo abrir el archivo.\n\n{ex}"); return
    
    # Detectar fecha de reunión desde el Word
    _MESES_WORD = {'ENERO','FEBRERO','MARZO','ABRIL','MAYO','JUNIO',
                   'JULIO','AGOSTO','SEPTIEMBRE','OCTUBRE','NOVIEMBRE','DICIEMBRE'}
    _pat_fecha_word = re.compile(
        r'(?:LUNES|MARTES|MI[EÉ]RCOLES|JUEVES|VIERNES|S[AÁ]BADO|DOMINGO)\s+(\d{1,2})\s+DE\s+(\w+)\s+DE\s+(\d{4})',
        re.IGNORECASE | re.UNICODE)
    for _para in doc.paragraphs:
        _m = _pat_fecha_word.search(_para.text.upper())
        if _m:
            _dia, _mes, _anio = _m.group(1), _m.group(2).upper(), _m.group(3)
            if _mes in _MESES_WORD:
                entry_fecha.delete(0, tk.END)
                entry_fecha.insert(0, f"{int(_dia):02d} DE {_mes} DE {_anio}")
                break

    # Auto-incrementar número de reunión
    _ultima_reunion = _leer_prefs().get("ultima_reunion", 0)
    if _ultima_reunion:
        entry_nro_reunion.delete(0, tk.END)
        entry_nro_reunion.insert(0, str(_ultima_reunion + 1))

    global DATOS_WORD_CACHED; DATOS_WORD_CACHED = []; curr = {}; capturing = False
    KEYWORDS = ("TURNO", "CLASICO", "CLÁSICO", "ESPECIAL", "HANDICAP", "GRAN PREMIO")
    for para in doc.paragraphs:
        txt = para.text.strip(); 
        if not txt: continue
        upper = txt.upper()
        if "LIQUIDARAN" in upper or "COMPUTAN" in upper or "INSCRIPCION" in upper: capturing = False; continue
        es_titulo = False
        if upper.startswith("PREMIO") and not upper.startswith("PREMIOS:"): es_titulo = True
        for k in KEYWORDS:
            if upper.startswith(k): es_titulo = True; break
        if es_titulo:
            if curr: DATOS_WORD_CACHED.append(curr)
            curr = {"nombre": txt, "distancia": "", "condicion_raw": "", "premios": ""}
            capturing = True
            m = re.search(r'(\d{1,2}[.]\d{3}|\d{3,4})\s*(?:m|mts|metros)', txt, re.I)
            if m: curr["distancia"] = m.group(1)
            continue
        if upper.startswith("PREMIOS:"): 
            if len(txt) > 120: continue 
            curr["premios"] = txt.split(':', 1)[1].strip(); capturing = False; continue
        if capturing and curr:
            if not curr["distancia"]:
                m = re.search(r'(\d{1,2}[.]\d{3}|\d{3,4})\s*(?:m|mts|metros)', txt, re.I)
                if m: curr["distancia"] = m.group(1)
                elif re.match(r'^\d{3,4}$', txt): curr["distancia"] = txt 
            if len(txt) > 10 and not re.match(r'^\d+$', txt): 
                if curr["condicion_raw"]: curr["condicion_raw"] += " " + txt
                else: curr["condicion_raw"] = txt
    if curr: DATOS_WORD_CACHED.append(curr)
    vals = [c.get("nombre", "Carrera") for c in DATOS_WORD_CACHED]; combo_word['values'] = vals
    if vals: combo_word.current(0); messagebox.showinfo("Cargado", f"{len(vals)} carreras detectadas."); aplicar_seleccion_word(None)
    else: messagebox.showwarning("Atención", "No se detectaron carreras.")

def aplicar_seleccion_word(e):
    idx = combo_word.current(); 
    if idx < 0: return
    d = DATOS_WORD_CACHED[idx]
    dist_orig = d.get("distancia",""); dist_key = dist_orig.replace('.', '').strip()
    entry_distancia.delete(0, tk.END)
    if dist_key in RECORDS: entry_distancia.insert(0, RECORDS[dist_key]); combo_dist.set(dist_key)
    else: entry_distancia.insert(0, dist_orig + " metros")
    entry_premios_dinero.delete(0, tk.END)
    try: dv = int(dist_key)
    except: dv = 0
    cat = "CAT. EXTRAOFICIAL" if dv <= 800 and dv > 0 else "CAT. INTERIOR"
    p_raw = d.get('premios','')
    if "COMPUTABLE" not in p_raw.upper(): entry_premios_dinero.insert(0, f"NO COMPUTABLE - {cat} - Premios: {p_raw}")
    else: entry_premios_dinero.insert(0, p_raw)
    entry_condicion.delete(0, tk.END); entry_condicion.insert(0, d.get("condicion_raw","").strip())

# =============================================================================
#  SECCIÓN 3: PERSISTENCIA
# =============================================================================

def accion_guardar_proyecto():
    if not programa_completo: return
    f = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("Proyecto Programa", "*.json")])
    if not f: return
    estado = {"fecha": entry_fecha.get(), "reunion": entry_nro_reunion.get(), "carreras": programa_completo}
    try:
        with open(f, 'w', encoding='utf-8') as json_file: json.dump(estado, json_file, indent=4)
        messagebox.showinfo("Guardado", "Proyecto guardado.")
    except Exception as e: messagebox.showerror("Error", str(e))

def accion_cargar_proyecto():
    f = filedialog.askopenfilename(filetypes=[("Proyecto Programa", "*.json")])
    if not f: return
    try:
        with open(f, 'r', encoding='utf-8') as json_file: estado = json.load(json_file)
        entry_fecha.delete(0, tk.END); entry_fecha.insert(0, estado.get("fecha", ""))
        entry_nro_reunion.delete(0, tk.END); entry_nro_reunion.insert(0, estado.get("reunion", "22"))
        global programa_completo; programa_completo = estado.get("carreras", [])
        _refrescar_lista_carreras(); limpiar_formulario(); messagebox.showinfo("Cargado", "Proyecto cargado.")
    except Exception as e: messagebox.showerror("Error", str(e))

# =============================================================================
#  SECCIÓN 4: AUXILIARES
# =============================================================================

def _clean_str(txt): return str(txt).replace('"', '').replace("Hs.", "").strip()
def _parse_money(txt):
    if not txt: return 0
    limpio = re.sub(r'[^\d]', '', str(txt))
    return int(limpio) if limpio else 0

# =============================================================================
#  SECCIÓN 5: EXCEL - PORTADA + CARRERAS + DISTRIBUCIÓN
# =============================================================================

def _estimar_filas_carrera(c):
    n_cond = max(1, len([x for x in c['cabecera']['condicion'].split('|') if x.strip()]))
    n_cab  = len(c['tabla_caballos'])
    n_acts = sum(1 for l in c['actuaciones'].split('\n') if l.strip())
    return 6 + n_cond + n_cab + n_acts

def _distribuir_carreras(carreras, n_pages):
    n = len(carreras)
    if n == 0: return []
    n_pages = max(1, min(n_pages, n // 2))
    base = n // n_pages; extra = n % n_pages
    grupos = []; i = 0
    for p in range(n_pages):
        sz = base + (1 if p < extra else 0)
        if sz > 0:
            grupos.append(carreras[i:i + sz]); i += sz
    return grupos

def _escribir_carrera_xl(ws, r, c, thin, med):
    cab = c['cabecera']
    # Fila: Nº carrera | Premio | Horario
    ws.row_dimensions[r].height = 21
    ws.merge_cells(f'C{r}:I{r}')
    ws[f'C{r}'].value = cab['premio'].upper()
    ws[f'C{r}'].font = Font(name='Tahoma', size=15, bold=True)
    ws[f'C{r}'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells(f'A{r}:B{r}')
    nro_c = ws.cell(row=r, column=1, value=f"{cab['nro_carrera']}º Carrera")
    nro_c.fill = PatternFill("solid", fgColor="000000")
    nro_c.font = Font(name='Arial Narrow', size=11, color="FFFFFF", bold=True)
    nro_c.alignment = Alignment(horizontal='center', vertical='center')
    hor_c = ws.cell(row=r, column=10, value=cab['horario'])
    hor_c.fill = PatternFill("solid", fgColor="000000")
    hor_c.font = Font(name='Arial Narrow', size=11, color="FFFFFF", bold=True)
    hor_c.alignment = Alignment(horizontal='center', vertical='center')
    r += 1
    # Distancia
    ws.row_dimensions[r].height = 15
    ws.merge_cells(f'A{r}:J{r}')
    ws.cell(row=r, column=1, value=cab['distancia']).alignment = Alignment(horizontal='center')
    ws.cell(row=r, column=1).font = Font(name='Utsaah', size=9, bold=True)
    r += 1
    # Condición (puede tener varios | )
    for lin in ([x.strip() for x in cab['condicion'].split('|') if x.strip()] or [""]):
        ws.row_dimensions[r].height = 15
        ws.merge_cells(f'A{r}:J{r}')
        ws.cell(row=r, column=1, value=lin).alignment = Alignment(wrap_text=True)
        ws.cell(row=r, column=1).font = Font(name='Utsaah', size=7)
        r += 1
    # Premios + Apuesta
    ws.row_dimensions[r].height = 15
    ws.merge_cells(f'A{r}:H{r}')
    ws.cell(row=r, column=1, value=cab['premios_dinero']).font = Font(name='Arial Narrow', size=8, bold=True)
    ws.merge_cells(f'I{r}:J{r}')
    ap_c = ws.cell(row=r, column=9, value=cab['apuesta'])
    ap_c.font = Font(name='Arial Black', size=9, bold=True, italic=True)
    ap_c.alignment = Alignment(horizontal='center', vertical='center')
    r += 1
    # Incremento
    ws.row_dimensions[r].height = 15
    ws.merge_cells(f'A{r}:H{r}')
    ws.cell(row=r, column=1, value=cab['incremento_2']).font = Font(name='Arial Narrow', size=8, bold=True)
    inc_val = _parse_money(cab['incremento'])
    if inc_val > 0:
        ws.merge_cells(f'I{r}:J{r}')
        ci = ws.cell(row=r, column=9, value=f"INCREMENTO: $ {inc_val:,.0f}".replace(",", "."))
        ci.font = Font(name='Arial Black', size=9, bold=True, italic=True)
        ci.alignment = Alignment(horizontal='center', vertical='center')
    r += 1
    # Headers tabla
    fila_inicio_tabla = r
    ws.row_dimensions[r].height = 15
    headers = ['4 Ult.', 'Nº', 'Caballo', 'Pelo', 'Jockey', 'E Kg', 'Padre-Madre', '', 'Caballeriza', 'Cuidador']
    ws.merge_cells(f'G{r}:H{r}')
    ws.cell(row=r, column=7).value = 'Padre - Madre'
    for col, h in enumerate(headers, 1):
        if col not in (7, 8):
            ws.cell(row=r, column=col, value=h).font = Font(name='Calibri', size=8, bold=True)
    r += 1
    # Caballos
    for row_data in c['tabla_caballos']:
        ws.row_dimensions[r].height = 12
        ws.merge_cells(f'G{r}:H{r}')
        for i in range(6): ws.cell(row=r, column=i + 1, value=row_data[i])
        ws.cell(row=r, column=7, value=row_data[6])
        ws.cell(row=r, column=9, value=row_data[7])
        ws.cell(row=r, column=10, value=row_data[8])
        r += 1
    fila_inicio_act = r
    # Actuaciones
    for l in c['actuaciones'].split('\n'):
        if not l.strip(): continue
        if "Debutante" not in l and " - " not in l[-5:]: l += " - PN"
        parts = l.split("||"); part1 = parts[0].strip(); rec = parts[1].strip() if len(parts) > 1 else ""
        m_a = re.match(r'^(\d+)[-\s]+(.*)', part1)
        num_x = int(m_a.group(1)) if m_a else 0; ant = m_a.group(2).strip() if m_a else part1
        ws.row_dimensions[r].height = 10
        ws.cell(row=r, column=1, value=num_x)
        ws.merge_cells(f'B{r}:F{r}'); ws.cell(row=r, column=2, value=ant)
        ws.merge_cells(f'G{r}:J{r}'); ws.cell(row=r, column=7, value=rec)
        r += 1
    fila_fin = r - 1
    # Bordes + alineación
    for row_cells in ws.iter_rows(min_row=fila_inicio_tabla, max_row=fila_fin, min_col=1, max_col=10):
        for cell in row_cells:
            b = Border(left=med, right=med, top=thin, bottom=thin)
            if cell.row == fila_inicio_tabla: b.top = med
            if cell.row == fila_fin: b.bottom = med
            if cell.column == 1: b.left = med
            if cell.column == 10: b.right = med
            if cell.row == fila_inicio_act - 1: b.bottom = med
            cell.border = b
            h_align = 'center'
            if cell.row >= fila_inicio_act:
                if cell.column in (2, 7): h_align = 'left'
                cell.font = Font(name='Calibri', size=7)
                if cell.column == 1: cell.font = Font(name='Calibri', size=8, bold=True)
            elif cell.row == fila_inicio_tabla:
                cell.font = Font(name='Calibri', size=8, bold=True)
            else:
                if cell.column == 1: h_align = 'right'
                elif cell.column in (3, 5, 7, 9, 10): h_align = 'left'
                cell.font = Font(name='Calibri', size=8, bold=(cell.column in (2, 3)))
            cell.alignment = Alignment(horizontal=h_align, vertical='center')
    r += 1  # separador entre carreras
    return r, fila_inicio_tabla, fila_inicio_act, fila_fin

def _make_rect_shape_elem(col0, row0, width_emu, height_emu, lines,
                          font_name, font_size_pt, bold, italic, shape_id, shape_name,
                          font_sizes=None, col_off=0, row_off=0):
    """Crea un elemento lxml oneCellAnchor con forma rectangular flotante (text box).
    font_sizes: lista opcional de tamaños por línea (sobrescribe font_size_pt por línea).
    """
    from lxml import etree
    XDR = 'http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing'
    A   = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    b_v = '1' if bold   else '0'
    i_v = '1' if italic else '0'
    paras = ''
    for idx, ln in enumerate(lines):
        sz = int((font_sizes[idx] if font_sizes and idx < len(font_sizes) else font_size_pt) * 100)
        lt = ln.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        paras += (
            f'<a:p xmlns:a="{A}"><a:pPr algn="ctr"/><a:r>'
            f'<a:rPr lang="es-AR" sz="{sz}" b="{b_v}" i="{i_v}" dirty="0">'
            f'<a:latin typeface="{font_name}"/></a:rPr>'
            f'<a:t>{lt}</a:t></a:r></a:p>'
        )
    xml = (
        f'<xdr:oneCellAnchor xmlns:xdr="{XDR}" xmlns:a="{A}">'
        f'<xdr:from><xdr:col>{col0}</xdr:col><xdr:colOff>{col_off}</xdr:colOff>'
        f'<xdr:row>{row0}</xdr:row><xdr:rowOff>{row_off}</xdr:rowOff></xdr:from>'
        f'<xdr:ext cx="{width_emu}" cy="{height_emu}"/>'
        f'<xdr:sp macro="" textlink=""><xdr:nvSpPr>'
        f'<xdr:cNvPr id="{shape_id}" name="{shape_name}"/>'
        f'<xdr:cNvSpPr txBox="1"/></xdr:nvSpPr>'
        f'<xdr:spPr><a:xfrm><a:off x="0" y="0"/>'
        f'<a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'<a:solidFill><a:srgbClr val="FFFFFF"/></a:solidFill>'
        f'<a:ln w="25400"><a:solidFill><a:srgbClr val="000000"/></a:solidFill></a:ln>'
        f'</xdr:spPr>'
        f'<xdr:txBody><a:bodyPr anchor="ctr" wrap="square" lIns="91440" rIns="91440" tIns="91440" bIns="91440"/><a:lstStyle/>'
        f'{paras}</xdr:txBody></xdr:sp><xdr:clientData/>'
        f'</xdr:oneCellAnchor>'
    )
    return etree.fromstring(xml)


def _inject_shapes_xlsx(fp, shape_elems):
    """Inyecta formas en drawing1.xml del XLSX ya guardado."""
    if not shape_elems:
        return
    import zipfile, io
    from lxml import etree
    buf = io.BytesIO(open(fp, 'rb').read())
    out = io.BytesIO()
    with zipfile.ZipFile(buf, 'r') as zin:
        names   = zin.namelist()
        drw_file = next((n for n in names
                         if n.startswith('xl/drawings/drawing') and n.endswith('.xml')), None)
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if drw_file and item.filename == drw_file:
                    root = etree.fromstring(data)
                    for elem in shape_elems:
                        root.append(elem)
                    data = etree.tostring(root, xml_declaration=True,
                                          encoding='UTF-8', standalone=True)
                zout.writestr(item, data)
    out.seek(0)
    with open(fp, 'wb') as f:
        f.write(out.read())


def _escribir_portada_xl(ws, fecha_txt, nro_reunion, shapes_out=None):
    thin = Side(style='thin'); med = Side(style='medium')

    def _borde(r1, c1, r2, c2, s=None):
        s = s or med
        for rr in range(r1, r2 + 1):
            for cc in range(c1, c2 + 1):
                cell = ws.cell(row=rr, column=cc)
                ex = cell.border
                cell.border = Border(
                    left=s   if cc == c1 else (ex.left   if ex else Side(style=None)),
                    right=s  if cc == c2 else (ex.right  if ex else Side(style=None)),
                    top=s    if rr == r1 else (ex.top    if ex else Side(style=None)),
                    bottom=s if rr == r2 else (ex.bottom if ex else Side(style=None)))

    # ── CABECERA 1: Logo + Título + Reunión (filas 1-2) ─────────────────────────
    ws.row_dimensions[1].height = 27.75
    ws.row_dimensions[2].height = 28.5
    ws.merge_cells('C1:J1')
    ws['C1'].value = "HIPÓDROMO DE TUCUMÁN - PROGRAMA OFICIAL"
    ws['C1'].font = Font(name='Arial Black', size=16, bold=True, italic=True)
    ws['C1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('C2:J2')
    ws['C2'].value = f"REUNIÓN Nº {nro_reunion} - {fecha_txt}"
    ws['C2'].font = Font(name='Arial Black', size=16, bold=True, italic=True)
    ws['C2'].alignment = Alignment(horizontal='center', vertical='center')
    if HAS_OXL_IMAGE:
        lp = ASSETS_DIR / "logo.png"
        if lp.exists():
            try:
                lg = OXLImage(str(lp)); lg.width = 67; lg.height = 68; lg.anchor = "A1"
                ws.add_image(lg)
            except Exception: pass
    _borde(1, 1, 2, 10)

    # ── CABECERA 2: Comisión (filas 3-6) ─────────────────────────────────────────
    ws.row_dimensions[3].height = 16
    for rng, txt in [('A3:E3', 'COMISIÓN DE CARRERAS'), ('F3:H3', 'VOCALES'), ('I3:J3', 'DELEGADO HIPODROMO')]:
        ws.merge_cells(rng); col = rng.split(':')[0]
        ws[col].value = txt
        ws[col].font = Font(name='Calibri', size=11, bold=True, underline='single')
        ws[col].alignment = Alignment(horizontal='center', vertical='center')
    for rn, rol, nom, vocal, deleg in [
        (4, "PRESIDENTE:",       "Dr. Luis Alberto Gamboa",       "Juan Ramon Rouges",  "Estanislao Perez Garcia"),
        (5, "VICE-PRESIDENTE:", "C.P.N Ernesto José Vidal Sanz", "Marcos Bruchmann",   ""),
        (6, "SECRETARIO:",      "Ignacio Lopez Bustos",           "Santiago Allende",   ""),
    ]:
        ws.row_dimensions[rn].height = 15
        ws.merge_cells(f'A{rn}:E{rn}')
        cell_rol = ws.cell(row=rn, column=1)
        try:
            from openpyxl.cell.rich_text import CellRichText, TextBlock
            from openpyxl.cell.text import InlineFont
            n_esp = max(2, 25 - len(rol))
            cell_rol.value = CellRichText(
                TextBlock(InlineFont(rFont='Calibri', sz=10, b=True, u='single'), rol),
                TextBlock(InlineFont(rFont='Calibri', sz=10), ' ' * n_esp + nom))
        except Exception:
            cell_rol.value = f"{rol}  {nom}"
            cell_rol.font = Font(name='Calibri', size=10, bold=True, underline='single')
        cell_rol.alignment = Alignment(horizontal='left', vertical='center')
        ws.merge_cells(f'F{rn}:H{rn}')
        ws.cell(row=rn, column=6, value=vocal).font = Font(name='Calibri', size=10)
        ws.cell(row=rn, column=6).alignment = Alignment(horizontal='center', vertical='center')
        if deleg:
            ws.merge_cells(f'I{rn}:J{rn}')
            ws.cell(row=rn, column=9, value=deleg).font = Font(name='Calibri', size=10)
            ws.cell(row=rn, column=9).alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[6].height = 15.75
    _borde(3, 1, 6, 10)

    # ── CABECERA 3: Admisión y Permanencia (fila 7) ──────────────────────────────
    ws.row_dimensions[7].height = 52
    ws.merge_cells('A7:J7')
    label = "Admisión y Permanencia: "
    body  = ("Las autoridades del Hipódromo de Tucumán ejercen la facultad de admisión y permanencia en las "
             "instalaciones del Hipódromo durante el desarrollo de la reunión hípica. Los profesionales y el "
             "público asistente se someten a las disposiciones del Reglamento General de Carreras y a las "
             "resoluciones de la Honorable Comisión de Carreras, cuyos fallos son inapelables. Los Boletos "
             "no cobrados solo se pagarán, los días de carreras de Tucumán y en el horario en que se desarrolle "
             "la reunión y tendrán validez, hasta 2 reuniones siguientes.-")
    try:
        from openpyxl.cell.rich_text import CellRichText, TextBlock
        from openpyxl.cell.text import InlineFont
        ws['A7'].value = CellRichText(
            TextBlock(InlineFont(rFont='Arial', sz=9, b=True, u='single'), label),
            TextBlock(InlineFont(rFont='Arial Narrow', sz=9), body))
    except Exception:
        ws['A7'].value = label + body
        ws['A7'].font = Font(name='Arial Narrow', size=9)
    ws['A7'].alignment = Alignment(wrap_text=True, horizontal='justify', vertical='top')
    _borde(7, 1, 7, 10)

    # ── Espacio + Carteles (filas 8-9) — formas flotantes ────────────────────────
    ws.row_dimensions[8].height = 4
    ws.row_dimensions[9].height = 54.75  # 1.93 cm
    # Formas rectangulares flotantes (8.8 cm × 1.93 cm cada una, con gap de ~0.5 cm)
    _W = 3168000; _H = 694800  # 8.8 cm × 1.93 cm en EMU
    _GAP = 180000              # 0.5 cm de separación entre ambas formas
    if shapes_out is not None:
        shapes_out.append(_make_rect_shape_elem(
            0, 8, _W, _H,
            ["El juego compulsivo, es perjudicial para la salud."],
            'Times New Roman', 14.5, True, True, 100, 'CartelJuego'))
        _W2 = 3416400  # 9.49 cm en EMU
        shapes_out.append(_make_rect_shape_elem(
            5, 8, _W2, _H,
            ["Los retirados en las apuestas combinadas", "(encadenadas), pasan al favorito."],
            'Times New Roman', 14.5, True, True, 101, 'CartelApuestas',
            col_off=_GAP))

    # ── Separador antes de 1ª carrera (fila 10) ─────────────────────────────────
    ws.row_dimensions[10].height = 4
    return 11

def _escribir_incrementos_xl(ws, row, programa, shapes_out=None):
    total_inc = 0; items = []
    for idx, car in enumerate(programa):
        cab = car['cabecera']
        monto = _parse_money(cab['incremento'])
        if monto <= 0: continue
        total_inc += monto
        nom_ap = re.sub(r'(?i)apuesta\s*', '', cab['apuesta']).strip().upper()
        m_entry = re.search(r'(\$\s*\d+[\.,]?\d*)', nom_ap)
        entry_str = m_entry.group(1).strip() if m_entry else ""
        type_name = re.sub(r'\$.*', '', nom_ap).strip()
        rango = 1
        for palabra, rv in [("QUINTUPLO", 5), ("CADENA", 6), ("CUATERNA", 4), ("TRIPLO", 3), ("DOBLE", 2)]:
            if palabra in nom_ap: rango = rv; break
        try: nro_start = int(cab['nro_carrera'])
        except: nro_start = idx + 1
        nros = [f"{nro_start + i}°" for i in range(rango)]
        if rango == 1:   c_str = f"{nros[0]} carrera"
        elif rango == 2: c_str = f"{nros[0]} y {nros[1]} carrera"
        else:            c_str = "; ".join(nros[:-1]) + f" y {nros[-1]} carrera"
        monto_fmt = f"$ {monto:,.0f}".replace(",", ".")
        items.append(f"{type_name} {entry_str}: {monto_fmt}.- {c_str}.-")
    if not items: return row
    ws.row_dimensions[row].height = 4; row += 1  # spacer
    # Forma flotante (19.1 cm ancho × variable alto)
    total_fmt = f"$ {total_inc:,.0f}".replace(",", ".")
    title_text = f"INCREMENTOS EN LA REUNION: {total_fmt}.-"
    height_cm  = 0.6 + 0.42 * (1 + len(items))  # título + items + padding interno
    height_pts = height_cm * 28.35
    height_emu = int(height_cm * 360000)
    width_emu  = 6876000  # 19.1 cm
    ws.row_dimensions[row].height = height_pts  # reserva espacio
    if shapes_out is not None:
        shape_row0 = row - 1  # 0-indexed row for the drawing anchor
        lines = [title_text] + items
        szs   = [11] + [9] * len(items)
        shapes_out.append(_make_rect_shape_elem(
            0, shape_row0, width_emu, height_emu,
            lines, 'Arial Black', 9, True, False, 102, 'Incrementos', font_sizes=szs))
    row += 1
    return row

def _escribir_logos_xl(ws, row):
    ws.row_dimensions[row].height = 5; row += 1
    if HAS_OXL_IMAGE:
        for path, anchor, iw, ih in [
            (ASSETS_DIR / "whatsapp.png", f"A{row}", 191, 28),
            (ASSETS_DIR / "redes.png",    f"H{row}", 227, 25),
        ]:
            if path.exists():
                try:
                    img = OXLImage(str(path)); img.width = iw; img.height = ih; img.anchor = anchor
                    ws.add_image(img)
                except Exception: pass
    ws.row_dimensions[row].height = 40; row += 1
    return row

def exportar_programa_excel():
    if not programa_completo: messagebox.showwarning("Vacío", "No hay datos."); return
    fecha_txt = (entry_fecha.get().strip() if entry_fecha and entry_fecha.get().strip()
                 else date.today().strftime("%d DE %B DE %Y").upper())
    nro_reunion = (entry_nro_reunion.get().strip() if entry_nro_reunion and entry_nro_reunion.get().strip() else "1")
    try: _guardar_prefs({"ultima_reunion": int(nro_reunion)})
    except: pass
    fp = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Workbook", "*.xlsx")],
                                      initialdir=str(PROGRAMAS_DIR), initialfile=f"{fecha_txt}.xlsx")
    if not fp: return
    try:
        wb = Workbook(); ws = wb.active; ws.title = "Programa"
        ws.page_margins = PageMargins(left=0.276, right=0.236, top=0.276, bottom=0.276,
                                      header=0.315, footer=0.315)
        ws.page_setup.paperSize = 9          # A4
        ws.page_setup.orientation = 'portrait'
        ws.print_options.horizontalCentered = True
        thin = Side(style="thin"); med = Side(style="medium")
        # Asegurar que exista un drawing para poder inyectar formas después
        from openpyxl.drawing.spreadsheet_drawing import SpreadsheetDrawing
        if ws._drawing is None:
            ws._drawing = SpreadsheetDrawing()
        shapes_out = []  # formas flotantes a inyectar post-guardado

        MAX_PAGE_PTS = 750  # puntos imprimibles por hoja A4 con esos márgenes
        PORTADA_PTS  = 210  # altura fija de filas 1-10 (header portada)

        # --- PORTADA (filas 1-10) → retorna 11 ---
        r = _escribir_portada_xl(ws, fecha_txt, nro_reunion, shapes_out)
        p1_data_start = r  # fila 11 = inicio 1ª carrera

        # --- 1RA CARRERA ---
        r, fi_tab_p1, _, fi_fin_p1 = _escribir_carrera_xl(ws, r, programa_completo[0], thin, med)

        # --- CUADRO DE INCREMENTOS (después de 1ª carrera) ---
        r = _escribir_incrementos_xl(ws, r, programa_completo, shapes_out)

        # --- LOGOS AL PIE ---
        r = _escribir_logos_xl(ws, r)

        # --- STRETCH página 1: escalar solo filas de datos (caballos + actuaciones) ---
        p1_data_end = r - 1  # incluye fila de logos para que el presupuesto sea correcto
        p1_scalable = set(range(fi_tab_p1 + 1, fi_fin_p1 + 1))
        if p1_data_end >= p1_data_start:
            p1_fixed = sum(ws.row_dimensions[f].height or 13
                           for f in range(p1_data_start, p1_data_end + 1) if f not in p1_scalable)
            p1_sc_pts = sum(ws.row_dimensions[f].height or 13 for f in p1_scalable)
            p1_avail = MAX_PAGE_PTS - PORTADA_PTS - p1_fixed
            if p1_sc_pts > 0 and p1_avail > 0:
                factor1 = p1_avail / p1_sc_pts
                for f in p1_scalable:
                    h = ws.row_dimensions[f].height or 13
                    ws.row_dimensions[f].height = max(7.5, min(24, h * factor1))

        # --- SALTO DE PÁGINA: fin de portada+1ª carrera ---
        ws.row_breaks.append(Break(id=r - 1))

        # --- CARRERAS RESTANTES distribuidas en hojas ---
        if len(programa_completo) > 1:
            restantes = programa_completo[1:]
            target_pages = 4 if len(programa_completo) <= 9 else 6
            grupos = _distribuir_carreras(restantes, target_pages - 1)
            for grupo in grupos:
                start_r = r
                scalable_rows = []
                for carrera in grupo:
                    r, fi_tab, _, fi_fin = _escribir_carrera_xl(ws, r, carrera, thin, med)
                    scalable_rows.extend(range(fi_tab + 1, fi_fin + 1))
                end_r = r - 2
                scalable_set = set(scalable_rows)
                fixed_pts = sum(ws.row_dimensions[f].height or 13
                                for f in range(start_r, end_r + 1) if f not in scalable_set)
                sc_pts = sum(ws.row_dimensions[f].height or 13 for f in scalable_rows)
                avail = MAX_PAGE_PTS - fixed_pts
                if sc_pts > 0 and avail > 0:
                    factor = avail / sc_pts
                    for f in scalable_rows:
                        h = ws.row_dimensions[f].height or 13
                        ws.row_dimensions[f].height = max(7.5, min(24, h * factor))
                ws.row_breaks.append(Break(id=r - 1))

        # --- ANCHOS DE COLUMNA ---
        for k, w in dict(A=6.1, B=3.9, C=15.6, D=5.1, E=13.7, F=3.7, G=9, H=12.4, I=14.6, J=13.7).items():
            ws.column_dimensions[k].width = w
        # Área de impresión explícita
        ws.print_area = f"A1:J{r - 1}"
        wb.save(fp)
        _inject_shapes_xlsx(fp, shapes_out)
        messagebox.showinfo("Listo", "Excel guardado correctamente.")
    except Exception as e:
        traceback.print_exc(); messagebox.showerror("Error Excel", str(e))

# =============================================================================
#  SECCIÓN 6: UI CALLBACKS
# =============================================================================

def editar_jockey(event):
    item_id = tabla_programa.identify_row(event.y); col_id = tabla_programa.identify_column(event.x)
    if not item_id: return
    if col_id == '#5': # Columna Jockey
        vals = list(tabla_programa.item(item_id, 'values')); old_jockey = vals[4]
        new_jockey = simpledialog.askstring("Editar Jockey", f"Jockey actual: {old_jockey}\nNuevo Jockey:")
        if new_jockey is not None: vals[4] = new_jockey.strip(); tabla_programa.item(item_id, values=vals)

def generar_programa_en_tabla():
    if db_caballos.empty: messagebox.showwarning("BD vacía", "No hay datos."); return
    for i in tabla_programa.get_children(): tabla_programa.delete(i)
    text_actuaciones.delete("1.0", tk.END)
    
    nombres = [n.strip() for n in text_caballos.get("1.0", tk.END).strip().split('\n') if n.strip()]
    kilos_manuales = [k.strip() for k in text_kilos.get("1.0", tk.END).strip().split('\n') if k.strip()]
    numero = 1; ultimo = None
    
    # El "enumerate" es clave acá para que la "i" vaya contando 0, 1, 2... y coincida con los kilos
    for i, nombre_in in enumerate(nombres):
        es_a = False; nombre = nombre_in
        if re.search(r'\(\s*a\s*\)\s*$', nombre_in, flags=re.I): es_a = True; nombre = re.sub(r'\(\s*a\s*\)\s*$', '', nombre_in, flags=re.I).strip()
        
        # --- MEJORA: "¿Quisiste decir?" ---
        nombre_upper = nombre.upper()
        if not db_caballos.empty and nombre_upper not in db_caballos['Caballo'].values:
            lista_nombres_db = db_caballos['Caballo'].astype(str).tolist()
            sugerencias = difflib.get_close_matches(nombre_upper, lista_nombres_db, n=1, cutoff=0.7)
            if not sugerencias:
                # Segunda pasada: normalizar Y↔I (confusión frecuente)
                _yi = lambda s: s.replace('Y', 'I')
                nombre_norm = _yi(nombre_upper)
                mapa_norm = {_yi(n): n for n in lista_nombres_db}
                if nombre_norm in mapa_norm:
                    sugerencias = [mapa_norm[nombre_norm]]
                else:
                    sug_norm = difflib.get_close_matches(nombre_norm, list(mapa_norm.keys()), n=1, cutoff=0.85)
                    if sug_norm: sugerencias = [mapa_norm[sug_norm[0]]]
            if sugerencias:
                sugerencia = sugerencias[0]
                resp = messagebox.askyesno("Posible error de tipeo", f"No se encontró a '{nombre_upper}' en la base de datos.\n\n¿Quisiste decir '{sugerencia}'?")
                if resp:
                    nombre = sugerencia
        # -----------------------------------
        
        datos = obtener_datos_caballo(nombre, db_caballos, db_actuaciones)
        
        # --- MEJORA: KILOS MANUALES ---
        kilo_asignado = kilos_manuales[i] if i < len(kilos_manuales) else datos.get('Peso', '')
        datos['E Kg'] = f"{datos.get('Edad', '')} {kilo_asignado}".strip()
        # ------------------------------
        
        if es_a and ultimo is not None: nro = f"{ultimo}a"
        else: nro = str(numero); ultimo = numero
        if not es_a: numero += 1
        
        datos['Jockey-Descargo'] = ""; datos['Nº'] = nro; datos['Caballo'] = nombre.upper()
        tabla_programa.insert('', tk.END, values=[datos.get(c, '') for c in ['4 Ult.','Nº','Caballo','Pelo','Jockey-Descargo','E Kg','Padre - Madre','Caballeriza','Cuidador']])
        
        acts = datos.get('actuaciones')
        lineas_locales = []
        
        # --- 1. Armamos los textos Locales ---
        if acts is not None and not acts.empty:
            for _, a in acts.iterrows():
                f_fmt = a['Fecha'].strftime('%d/%m/%y') if pd.notna(a['Fecha']) else ''
                if str(a['Puesto Final']).strip().upper() == 'NC': 
                    lineas_locales.append(f"{f_fmt} - No Corrió.")
                    continue
                jk_full = str(a.get('Jockey',''))
                jk = f"{jk_full.split()[1][:1]}. {jk_full.split()[0]}" if len(jk_full.split())>1 else jk_full
                dist_txt = " (Dist.)" if str(a.get('Puesto Original')) != str(a.get('Puesto Final')) else ""
                pista = a.get('Pista', 'PN')
                if not pista: pista = 'PN'
                
                if str(a.get('Puesto Original')).strip() in ['1','1.0']:
                    margen = formatear_cuerpos(a.get('Margen',''))
                    lineas_locales.append(f"{f_fmt} - {jk} - 1º gan x {margen} a {str(a.get('Segundo','')).title()} - {a.get('Tiempo Ganador','')}{dist_txt} - {pista}")
                else:
                    cuerpos = formatear_cuerpos(a.get('Cuerpos al Ganador',''))
                    lineas_locales.append(f"{f_fmt} - {jk} - {a.get('Puesto Original')}º a {cuerpos} de {str(a.get('Ganador','')).title()} - {a.get('Tiempo Ganador','')}{dist_txt} - {pista}")

        # --- 2. Traemos los textos Externos ---
        texto_ext = datos.get('texto_act_ext', '')
        lineas_ext_raw = [l.strip() for l in texto_ext.split('\n') if l.strip() and l.strip().lower() != 'nan']

        # Filtramos líneas externas que sean duplicado de una actuación local en DB
        # (misma fecha DD/MM/YY → es la misma carrera local, ya está en lineas_locales)
        db_fechas = datos.get('todas_act_fechas', set())
        _dp = re.compile(r'^(\d{2}/\d{2}/\d{2})')
        lineas_ext = []
        for ln in lineas_ext_raw:
            m = _dp.match(ln)
            if m and m.group(1) in db_fechas:
                continue  # duplicado local, ignorar
            lineas_ext.append(ln)

        # Merge cronológico: ordenamos por fecha extraída del texto
        def _fecha_linea(ln):
            m = _dp.match(ln)
            if m:
                try: return datetime.strptime(m.group(1), '%d/%m/%y')
                except: pass
            return datetime.min

        todas_las_lineas = sorted(lineas_ext + lineas_locales, key=_fecha_linea)
        todas_las_lineas = todas_las_lineas[-2:] if todas_las_lineas else []
        
        if not todas_las_lineas:
            bloque = "Debutante"
        else:
            # Las unimos en el orden en el que están (vieja izquierda || nueva derecha)
            bloque = "   ||   ".join(todas_las_lineas)
            
        text_actuaciones.insert(tk.END, f"{nro}  {bloque}\n")
    _generar_detalle_apuestas()

def obtener_datos_formulario():
    rows = [tabla_programa.item(i)['values'] for i in tabla_programa.get_children()]
    return {"cabecera": {"nro_carrera": entry_nro_carrera.get(), "premio": entry_premio.get(), "horario": entry_horario.get(), "distancia": entry_distancia.get(), "condicion": entry_condicion.get(), "premios_dinero": entry_premios_dinero.get(), "apuesta": entry_apuesta.get(), "incremento": entry_incremento.get(), "incremento_2": entry_incremento_2.get()}, "tabla_caballos": rows, "actuaciones": text_actuaciones.get("1.0", tk.END).strip()}

def _actualizar_combo_apuesta():
    if entry_apuesta is None: return
    sesion   = [c['cabecera']['apuesta'] for c in programa_completo if c['cabecera'].get('apuesta','').strip()]
    historico = _leer_prefs().get('apuestas_usadas', [])
    vistos = set(); sugerencias = []
    for a in sesion + historico:
        if a and a not in vistos: sugerencias.append(a); vistos.add(a)
    entry_apuesta['values'] = sugerencias

def _generar_detalle_apuestas(event=None):
    if entry_apuesta is None or entry_incremento_2 is None: return
    apuesta = entry_apuesta.get().strip()
    n_caballos = len(tabla_programa.get_children()) if tabla_programa else 0
    exacta   = "IMPERFECTA $ 200"  if n_caballos > 13 else "EXACTA $ 200"
    trifecta = "CUATRIFECTA $ 200" if n_caballos > 13 else "TRIFECTA $ 200"
    apuesta_corta = re.sub(r'^APUESTA\s+', '', apuesta, flags=re.I).strip()
    apuesta_es_doble = bool(re.search(r'DOBLE', apuesta_corta, re.I))

    # Detectar última carrera: si ya hay una DOBLE FINAL guardada y el nro actual es mayor
    es_ultima = False
    try:
        nro_actual = int(entry_nro_carrera.get().strip())
        nros_doble_final = [int(c['cabecera']['nro_carrera']) for c in programa_completo
                            if 'DOBLE FINAL' in c['cabecera'].get('apuesta','').upper()
                            and str(c['cabecera']['nro_carrera']).isdigit()]
        if nros_doble_final and nro_actual > max(nros_doble_final):
            es_ultima = True
    except: pass

    partes = ["AP. A GANADOR $ 1 ($ 500 MINIMO)", exacta, trifecta]
    if not apuesta_es_doble and not es_ultima:
        partes.append("DOBLE $ 200")
    if apuesta_corta: partes.append(apuesta_corta)
    entry_incremento_2.delete(0, tk.END)
    entry_incremento_2.insert(0, ", ".join(partes))

def _on_apuesta_changed(event=None):
    if entry_apuesta is None or entry_incremento is None: return
    apuesta = entry_apuesta.get().strip()
    if not apuesta: return
    valor = _leer_prefs().get('incrementos', {}).get(apuesta, '')
    if valor:
        entry_incremento.delete(0, tk.END)
        entry_incremento.insert(0, valor)
    _generar_detalle_apuestas()

def limpiar_formulario():
    for e in [entry_nro_carrera, entry_premio, entry_horario, entry_distancia, entry_condicion, entry_premios_dinero, entry_incremento, entry_incremento_2]: e.delete(0, tk.END)
    if entry_apuesta is not None: entry_apuesta.set('')
    text_caballos.delete("1.0", tk.END); text_kilos.delete("1.0", tk.END); text_actuaciones.delete("1.0", tk.END)
    for i in tabla_programa.get_children(): tabla_programa.delete(i)
    global indice_edicion; indice_edicion = None; btn_accion.config(text="Añadir Carrera")
    _actualizar_combo_apuesta()
    if not programa_completo:
        hist = _leer_prefs().get('apuestas_usadas', [])
        entry_apuesta.set(hist[0] if hist else "APUESTA CUATERNA $ 300")

def guardar_o_anadir_carrera():
    if not tabla_programa.get_children(): return
    data = obtener_datos_formulario(); global indice_edicion
    if indice_edicion is not None: programa_completo[indice_edicion] = data; messagebox.showinfo("OK", "Carrera Actualizada")
    else: programa_completo.append(data); messagebox.showinfo("OK", "Carrera Añadida")
    apuesta_usada = data['cabecera'].get('apuesta', '').strip()
    incremento_usado = data['cabecera'].get('incremento', '').strip()
    if apuesta_usada:
        hist = _leer_prefs().get('apuestas_usadas', [])
        if apuesta_usada not in hist: hist.insert(0, apuesta_usada); _guardar_prefs({'apuestas_usadas': hist[:20]})
    if apuesta_usada and incremento_usado:
        incs = _leer_prefs().get('incrementos', {})
        incs[apuesta_usada] = incremento_usado
        _guardar_prefs({'incrementos': incs})
    _refrescar_lista_carreras(); limpiar_formulario()

def cargar_carrera_para_editar():
    sel = lista_carreras.curselection(); 
    if not sel: return
    idx = int(sel[0]); limpiar_formulario(); global indice_edicion; indice_edicion = idx; btn_accion.config(text="Guardar Cambios")
    data = programa_completo[idx]; cab = data['cabecera']
    entry_nro_carrera.insert(0, cab['nro_carrera']); entry_premio.insert(0, cab['premio']); entry_horario.insert(0, cab['horario']); entry_distancia.insert(0, cab['distancia']); entry_condicion.insert(0, cab['condicion']); entry_premios_dinero.insert(0, cab['premios_dinero']); entry_apuesta.insert(0, cab['apuesta']); entry_incremento.insert(0, cab['incremento']); entry_incremento_2.insert(0, cab['incremento_2'])
    for row in data['tabla_caballos']: tabla_programa.insert('', tk.END, values=row); text_caballos.insert(tk.END, row[2] + "\n"); ekg_partes = str(row[5]).strip().split(maxsplit=1)
    kilo_val = ekg_partes[1] if len(ekg_partes) > 1 else ""
    text_kilos.insert(tk.END, kilo_val + "\n")
    text_actuaciones.insert(tk.END, data['actuaciones']); messagebox.showinfo("Editando", f"Editando carrera {cab['nro_carrera']}.")

def eliminar_carrera():
    sel = lista_carreras.curselection(); 
    if not sel: return
    programa_completo.pop(int(sel[0])); _refrescar_lista_carreras(); limpiar_formulario()

def _refrescar_lista_carreras():
    lista_carreras.delete(0, tk.END)
    for c in programa_completo: lista_carreras.insert(tk.END, f"{c['cabecera']['nro_carrera']}º - {c['cabecera']['premio']}")
    contador_carreras.set(f"Carreras: {len(programa_completo)}")

def accion_reset_db():
    if messagebox.askyesno("Reset", "Borrar DB?"):
        if os.path.exists(NOMBRE_BD): os.remove(NOMBRE_BD)
        _inicializar_db_si_no_existe(); global db_caballos, db_actuaciones; db_caballos, db_actuaciones = conectar_y_cargar_datos(); messagebox.showinfo("Info", "DB Reiniciada")

def _fecha_de_programa(path: str):
    MESES = {'ENERO':'01','FEBRERO':'02','MARZO':'03','ABRIL':'04',
              'MAYO':'05','JUNIO':'06','JULIO':'07','AGOSTO':'08',
              'SEPTIEMBRE':'09','OCTUBRE':'10','NOVIEMBRE':'11','DICIEMBRE':'12'}
    stem = Path(path).stem.upper()
    m = re.match(r'(\d{1,2})\s+DE\s+(\w+)\s+DE\s+(\d{4})', stem)
    if m:
        dia, mes_txt, anio = int(m.group(1)), m.group(2), m.group(3)
        mes_n = MESES.get(mes_txt)
        if mes_n:
            return f"{dia:02d}-{mes_n}-{anio[-2:]}"
    return None

def _fecha_de_resultados(path: str):
    stem = Path(path).stem
    m = re.match(r'resultados?\s+(\d{2}-\d{2}-\d{2})', stem, re.I)
    return m.group(1) if m else None

def accion_importar_programa():
    try:
        from migracion import cargar_base_de_datos_caballos
    except ImportError as e:
        messagebox.showerror("Error de importación", f"No se pudo cargar migracion.py:\n{e}"); return

    fp = filedialog.askopenfilename(
        title="Seleccionar Excel de Programa",
        initialdir=str(PROGRAMAS_DIR),
        filetypes=[("Excel", "*.xlsx *.xls"), ("Todos", "*.*")]
    )
    if not fp: return

    fecha = _fecha_de_programa(fp)
    if not fecha:
        messagebox.showerror("Error de nombre",
            f"No se pudo determinar la fecha del archivo:\n{Path(fp).name}\n\n"
            "El nombre debe ser como: '12 DE ABRIL DE 2026.xlsx'")
        return

    _inicializar_db_si_no_existe()
    df_cab = cargar_base_de_datos_caballos(fp, fecha)
    if df_cab is None or df_cab.empty:
        messagebox.showwarning("Sin datos", "No se encontraron caballos en el archivo."); return

    conn = sqlite3.connect(NOMBRE_BD); cargados = 0
    for _, row in df_cab.iterrows():
        nombre = re.sub(r'\s+', ' ', str(row.get('CABALLO', ''))).strip().upper()
        if not nombre or nombre == 'NAN': continue
        e_kg = str(row.get('E Kg', '')).strip()
        m2 = re.match(r'^\D*(\d{1,2})\s+(\d{2,3})\D*$', e_kg)
        edad, peso = (m2.group(1), m2.group(2)) if m2 else ('', '')
        actu_ext = row.get('4 Ult.', '') or ''
        if 'debuta' in str(actu_ext).lower(): actu_ext = ''
        texto_ext = str(row.get('TEXTO_ACT_EXT_FINAL', '')).strip()
        conn.execute('INSERT OR IGNORE INTO caballos (nombre) VALUES (?)', (nombre,))
        conn.execute('''UPDATE caballos SET padre_madre=?, pelo=?, ultima_edad=?, ultimo_peso=?,
            ultimo_jockey=?, caballeriza=?, cuidador=?, ultima_actuacion_externa=?,
            texto_actuaciones_externas=?, snapshot_programa_fecha=? WHERE nombre=?''',
            (row.get('Padre - Madre',''), row.get('Pelo',''), edad, peso,
             row.get('Jockey-Descargo',''), row.get('Caballeriza',''),
             row.get('Cuidador',''), actu_ext, texto_ext, fecha, nombre))
        cargados += 1
    conn.commit(); conn.close()

    global db_caballos, db_actuaciones
    db_caballos, db_actuaciones = conectar_y_cargar_datos()
    _guardar_en_registro("programas", fecha, Path(fp).name)
    _actualizar_estado_db()
    messagebox.showinfo("Programa cargado", f"Fecha: {fecha}\nCaballos procesados: {cargados}")

def accion_importar_resultados():
    try:
        from migracion import cargar_historial_actuaciones
    except ImportError as e:
        messagebox.showerror("Error de importación", f"No se pudo cargar migracion.py:\n{e}"); return

    fp = filedialog.askopenfilename(
        title="Seleccionar Excel de Resultados",
        initialdir=str(RESULTADOS_DIR),
        filetypes=[("Excel", "*.xlsx *.xls"), ("Todos", "*.*")]
    )
    if not fp: return

    fecha = _fecha_de_resultados(fp)
    if not fecha:
        messagebox.showerror("Error de nombre",
            f"No se pudo determinar la fecha del archivo:\n{Path(fp).name}\n\n"
            "El nombre debe ser como: 'Resultados 12-04-26.xlsx'")
        return

    from datetime import datetime as _dt
    try:
        dt = _dt.strptime(fecha, '%d-%m-%y')
        fecha_iso_prefix = dt.strftime('%Y-%m-%d')
    except ValueError:
        fecha_iso_prefix = fecha

    _inicializar_db_si_no_existe()
    conn = sqlite3.connect(NOMBRE_BD); c = conn.cursor()
    ya_existe = c.execute("SELECT COUNT(*) FROM actuaciones WHERE fecha LIKE ?",
                          (f'{fecha_iso_prefix}%',)).fetchone()[0] > 0
    conn.close()

    if ya_existe:
        if not messagebox.askyesno("Resultados existentes",
            f"Ya existen resultados para la fecha {fecha}.\n¿Reemplazar?"): return
        conn = sqlite3.connect(NOMBRE_BD)
        conn.execute("DELETE FROM actuaciones WHERE fecha LIKE ?", (f'{fecha_iso_prefix}%',))
        conn.commit(); conn.close()

    df_act = cargar_historial_actuaciones(fp, fecha)
    if df_act is None or df_act.empty:
        messagebox.showwarning("Sin datos", "No se encontraron actuaciones en el archivo."); return

    colmap = {'Fecha':'fecha','CABALLO':'nombre_caballo','Puesto Original':'puesto_original',
              'Puesto Final':'puesto_final','Jockey':'jockey','Cuerpos al Ganador':'cuerpos',
              'Ganador':'ganador','Segundo':'segundo','Margen':'margen',
              'Tiempo Ganador':'tiempo_ganador','Pista':'pista',
              'Fue Distanciado':'fue_distanciado','Observacion':'observacion'}
    conn = sqlite3.connect(NOMBRE_BD)
    df_act.rename(columns=colmap).to_sql('actuaciones', conn, if_exists='append', index=False)
    tot = len(df_act); conn.close()

    global db_caballos, db_actuaciones
    db_caballos, db_actuaciones = conectar_y_cargar_datos()
    _guardar_en_registro("resultados", fecha, Path(fp).name)
    _actualizar_estado_db()
    messagebox.showinfo("Resultados cargados", f"Fecha: {fecha}\nActuaciones cargadas: {tot}")

# =============================================================================
#  SECCIÓN 7: STARTUP
# =============================================================================

db_caballos, db_actuaciones = conectar_y_cargar_datos()
root = tk.Tk(); root.title("Gestión de Programas Hípicos"); root.configure(bg=COLORS["bg"])
try: root.iconbitmap(str(ASSETS_DIR/"programa.ico"))
except: pass

# ---------- ESTILOS ttk ----------
style = ttk.Style(root)
style.theme_use('clam')
C = COLORS
style.configure("TFrame",      background=C["bg"])
style.configure("Card.TFrame", background=C["card"])
style.configure("TLabel",       background=C["bg"],   foreground=C["ink"], font=("Segoe UI", 9))
style.configure("Card.TLabel",  background=C["card"], foreground=C["ink"], font=("Segoe UI", 9))
style.configure("Field.TLabel", background=C["card"], foreground=C["ink"], font=("Segoe UI", 9))
style.configure("TButton",  font=("Segoe UI", 9), padding=(8, 4), relief="flat")
style.configure("Primary.TButton", background=C["primary"], foreground="white",
                font=("Segoe UI", 9), padding=(8, 4))
style.map("Primary.TButton", background=[("active", "#1c6e71"), ("pressed", "#155558")])
style.configure("Accent.TButton", background=C["accent"], foreground="white",
                font=("Segoe UI", 9, "bold"), padding=(8, 4))
style.map("Accent.TButton", background=[("active", "#d4502a"), ("pressed", "#b84020")])
style.configure("Green.TButton", background="#2a7d32", foreground="white",
                font=("Segoe UI", 9, "bold"), padding=(8, 4))
style.map("Green.TButton", background=[("active", "#236129"), ("pressed", "#1a4a1e")])
style.configure("Word.TButton", background="#1565c0", foreground="white",
                font=("Segoe UI", 9), padding=(8, 4))
style.map("Word.TButton", background=[("active", "#0d4f9e"), ("pressed", "#0a3a7a")])
style.configure("Danger.TButton", background="#b71c1c", foreground="white",
                font=("Segoe UI", 9), padding=(8, 4))
style.map("Danger.TButton", background=[("active", "#8e1515"), ("pressed", "#6b0f0f")])
style.configure("TLabelframe",       background=C["card"], bordercolor=C["line"])
style.configure("TLabelframe.Label", background=C["card"], foreground=C["primary"],
                font=("Segoe UI", 9, "bold"))
style.configure("TEntry",    padding=3)
style.configure("TCombobox", padding=3)
style.configure("Treeview",
    background=C["card"], fieldbackground=C["card"],
    foreground=C["ink"], rowheight=22, font=("Segoe UI", 9))
style.configure("Treeview.Heading",
    background=C["primary"], foreground="white",
    font=("Segoe UI", 9, "bold"), relief="flat")
style.map("Treeview.Heading", background=[("active", "#1c6e71")])
style.map("Treeview",
    background=[("selected", C["primary"])],
    foreground=[("selected", "white")])
style.configure("TSeparator", background=C["line"])

class ScrollableFrame(ttk.Frame):
    def __init__(self, container, *args, **kwargs):
        super().__init__(container, *args, **kwargs)
        canvas = tk.Canvas(self, bg=COLORS["bg"], highlightthickness=0)
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        self.scrollable_window = ttk.Frame(canvas, style="Card.TFrame")
        self.scrollable_window.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        self.window_id = canvas.create_window((0, 0), window=self.scrollable_window, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set); canvas.pack(side="left", fill="both", expand=True); scrollbar.pack(side="right", fill="y")
        canvas.bind_all("<MouseWheel>", lambda event: canvas.yview_scroll(int(-1*(event.delta/120)), "units")); canvas.bind("<Configure>", lambda e: canvas.itemconfig(self.window_id, width=e.width))

head = tk.Frame(root, bg=COLORS["primary"], padx=15, pady=10); head.pack(side=tk.TOP, fill=tk.X)
tk.Label(head, text="HIPÓDROMO DE TUCUMÁN  —  GESTIÓN DE PROGRAMAS",
         bg=COLORS["primary"], fg="white", font=("Segoe UI", 14, "bold")).pack(side=tk.LEFT)
ttk.Button(head, text="Cargar Carta de Llamada (.docx)", command=cargar_word_entrada,
           style="Word.TButton").pack(side=tk.RIGHT, padx=(5, 0))
ttk.Separator(root, orient='horizontal').pack(fill=tk.X)

ttk.Separator(root, orient='horizontal').pack(side=tk.BOTTOM, fill=tk.X)
foot = ttk.Frame(root, style="Card.TFrame", padding=(15, 10)); foot.pack(side=tk.BOTTOM, fill=tk.X)
contador_carreras = tk.StringVar(value="Carreras: 0")
ttk.Label(foot, textvariable=contador_carreras, style="Card.TLabel", font=("Segoe UI", 10, "bold"),
          foreground=COLORS["primary"]).pack(side=tk.LEFT)
estado_db_var = tk.StringVar(value=_leer_estado_db())
ttk.Label(foot, textvariable=estado_db_var, style="Card.TLabel",
          font=("Segoe UI", 8), foreground="#6b7280").pack(side=tk.LEFT, padx=25)
ttk.Button(foot, text="Exportar Excel", command=exportar_programa_excel, style="Green.TButton").pack(side=tk.RIGHT, padx=(5, 0))

main_scroll = ScrollableFrame(root); main_scroll.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
main_content = main_scroll.scrollable_window

form_container = ttk.Frame(main_content, style="Card.TFrame", padding=10); form_container.pack(fill=tk.X)
f1 = ttk.LabelFrame(form_container, text="Información de Carrera", padding=15); f1.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10)

ttk.Label(f1, text="Fecha Reunión (PDF):", style="Field.TLabel", foreground="red").grid(row=0, column=0, sticky="w", pady=5)
entry_fecha = ttk.Entry(f1); entry_fecha.grid(row=0, column=1, columnspan=3, sticky="we", pady=5)
entry_fecha.insert(0, "")  # Se completa automáticamente al cargar el Word

# --- NUEVO CAMPO REUNION ---
ttk.Label(f1, text="Nº Reunión (PDF):", style="Field.TLabel", foreground="blue").grid(row=1, column=0, sticky="w", pady=5)
entry_nro_reunion = ttk.Entry(f1, width=10); entry_nro_reunion.grid(row=1, column=1, sticky="w", pady=5)
_nro_ini = _leer_prefs().get("ultima_reunion", ""); entry_nro_reunion.insert(0, str(_nro_ini) if _nro_ini else "") 

ttk.Label(f1, text="Selección Word:", style="Field.TLabel").grid(row=2, column=0, sticky="w", pady=5)
combo_word = ttk.Combobox(f1, width=35, state="readonly"); combo_word.grid(row=2, column=1, columnspan=3, sticky="we", pady=5); combo_word.bind("<<ComboboxSelected>>", aplicar_seleccion_word)

ttk.Label(f1, text="Nº Carrera:", style="Field.TLabel").grid(row=3, column=0, sticky="w", pady=5); entry_nro_carrera = ttk.Entry(f1, width=10); entry_nro_carrera.grid(row=3, column=1, sticky="w", pady=5, padx=5)
ttk.Label(f1, text="Horario:", style="Field.TLabel").grid(row=3, column=2, sticky="w", pady=5); entry_horario = ttk.Entry(f1, width=15); entry_horario.grid(row=3, column=3, sticky="w", pady=5, padx=5)
ttk.Label(f1, text="Premio:", style="Field.TLabel").grid(row=4, column=0, sticky="w", pady=5); entry_premio = ttk.Entry(f1); entry_premio.grid(row=4, column=1, columnspan=3, sticky="we", pady=5)
ttk.Label(f1, text="Distancia:", style="Field.TLabel").grid(row=5, column=0, sticky="w", pady=5)
dist_var = tk.StringVar()
def _on_dist(*_): 
    k = dist_var.get().strip()
    if k in RECORDS: entry_distancia.delete(0, tk.END); entry_distancia.insert(0, RECORDS[k])
combo_dist = ttk.Combobox(f1, width=8, values=list(RECORDS.keys()), textvariable=dist_var); combo_dist.bind("<<ComboboxSelected>>", _on_dist); combo_dist.grid(row=5, column=1, sticky="w", pady=5, padx=5); entry_distancia = ttk.Entry(f1); entry_distancia.grid(row=5, column=2, columnspan=2, sticky="we", pady=5)
ttk.Label(f1, text="Condición (Usar | para salto):", style="Field.TLabel").grid(row=6, column=0, sticky="w", pady=5); entry_condicion = ttk.Entry(f1); entry_condicion.grid(row=6, column=1, columnspan=3, sticky="we", pady=5)
ttk.Label(f1, text="Premios ($):", style="Field.TLabel").grid(row=7, column=0, sticky="w", pady=5); entry_premios_dinero = ttk.Entry(f1); entry_premios_dinero.grid(row=7, column=1, columnspan=3, sticky="we", pady=5)

f2 = ttk.LabelFrame(form_container, text="Apuestas y Caballos", padding=15); f2.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10)
ttk.Label(f2, text="Apuesta (Título):", style="Field.TLabel").grid(row=0, column=0, sticky="w", pady=5)
entry_apuesta = ttk.Combobox(f2); entry_apuesta.grid(row=0, column=1, sticky="we", pady=5, padx=5)
entry_apuesta.bind('<<ComboboxSelected>>', _on_apuesta_changed)
entry_apuesta.bind('<Return>', _on_apuesta_changed)
_actualizar_combo_apuesta()
if not programa_completo:
    _hist0 = _leer_prefs().get('apuestas_usadas', [])
    entry_apuesta.set(_hist0[0] if _hist0 else "APUESTA CUATERNA $ 300")
ttk.Label(f2, text="Incremento ($):", style="Field.TLabel").grid(row=1, column=0, sticky="w", pady=5); entry_incremento = ttk.Entry(f2, width=15); entry_incremento.grid(row=1, column=1, sticky="w", pady=5, padx=5)
ttk.Label(f2, text="Detalle Apuestas:", style="Field.TLabel").grid(row=2, column=0, sticky="w", pady=5); entry_incremento_2 = ttk.Entry(f2, width=15); entry_incremento_2.grid(row=2, column=1, sticky="w", pady=5, padx=5)
ttk.Label(f2, text="Pegar Lista Caballos:").grid(row=3, column=0, sticky="nw", pady=5)
text_caballos = tk.Text(f2, height=6, width=22, bg=COLORS["card"], fg=COLORS["ink"],
    insertbackground=COLORS["primary"], font=("Segoe UI", 9), relief="solid", borderwidth=1)
text_caballos.grid(row=3, column=1, rowspan=3, sticky="we", pady=5, padx=5)

ttk.Label(f2, text="Kilos:", style="Field.TLabel").grid(row=3, column=2, sticky="nw", pady=5)
text_kilos = tk.Text(f2, height=6, width=8, bg=COLORS["card"], fg=COLORS["ink"],
    insertbackground=COLORS["primary"], font=("Segoe UI", 9), relief="solid", borderwidth=1)
text_kilos.grid(row=3, column=3, rowspan=3, sticky="we", pady=5, padx=5)

ttk.Separator(main_content, orient='horizontal').pack(fill=tk.X, padx=10)
btn_box = ttk.Frame(main_content, style="Card.TFrame", padding=(15, 10)); btn_box.pack(fill=tk.X)
ttk.Button(btn_box, text="Procesar Tabla", command=generar_programa_en_tabla, style="Primary.TButton").pack(side=tk.LEFT, padx=(0, 8))
btn_accion = ttk.Button(btn_box, text="Añadir Carrera", command=guardar_o_anadir_carrera, style="Accent.TButton")
btn_accion.pack(side=tk.LEFT, padx=(0, 8))
ttk.Button(btn_box, text="Limpiar Formulario", command=limpiar_formulario).pack(side=tk.LEFT)
ttk.Separator(main_content, orient='horizontal').pack(fill=tk.X, padx=10)

paned = ttk.PanedWindow(main_content, orient=tk.HORIZONTAL); paned.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
frame_left = ttk.Frame(paned, style="Card.TFrame"); paned.add(frame_left, weight=4)
cols = ['4 Ult.','Nº','Caballo','Pelo','Jockey-Descargo','E Kg','Padre - Madre','Caballeriza','Cuidador']
tabla_programa = ttk.Treeview(frame_left, columns=cols, show='headings', height=10)
col_widths = {'4 Ult.':55,'Nº':35,'Caballo':140,'Pelo':50,'Jockey-Descargo':115,'E Kg':50,'Padre - Madre':115,'Caballeriza':105,'Cuidador':95}
col_align = {'Caballo':'w','Jockey-Descargo':'w','Padre - Madre':'w','Caballeriza':'w','Cuidador':'w'}
for c in cols:
    tabla_programa.heading(c, text=c)
    tabla_programa.column(c, width=col_widths.get(c, 90), anchor=col_align.get(c, 'center'))
tabla_programa.pack(fill=tk.BOTH, expand=True, pady=(0, 6))
tabla_programa.bind("<Double-1>", editar_jockey)

frame_acts = ttk.Frame(frame_left, style="Card.TFrame"); frame_acts.pack(fill=tk.BOTH, expand=True)
ttk.Label(frame_acts, text="Actuaciones Generadas (Editable):", style="Field.TLabel").pack(anchor="w", pady=(4, 2))
scroll_acts = ttk.Scrollbar(frame_acts); scroll_acts.pack(side=tk.RIGHT, fill=tk.Y)
text_actuaciones = tk.Text(frame_acts, height=12, yscrollcommand=scroll_acts.set,
    bg=COLORS["card"], fg=COLORS["ink"], insertbackground=COLORS["primary"],
    font=("Consolas", 8), relief="solid", borderwidth=1)
text_actuaciones.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
scroll_acts.config(command=text_actuaciones.yview)

frame_right = ttk.Frame(paned, style="Card.TFrame"); paned.add(frame_right, weight=1)
ttk.Label(frame_right, text="Carreras en Programa:", style="Field.TLabel").pack(anchor="w", pady=(5, 2))
lista_carreras = tk.Listbox(frame_right, bg=COLORS["card"], fg=COLORS["ink"],
    selectbackground=COLORS["primary"], selectforeground="white",
    font=("Segoe UI", 9), borderwidth=1, relief="solid", activestyle="none")
lista_carreras.pack(fill=tk.BOTH, expand=True, padx=5)
ttk.Button(frame_right, text="Cargar para Editar", command=cargar_carrera_para_editar, style="Primary.TButton").pack(fill=tk.X, pady=(6, 3), padx=5)
ttk.Button(frame_right, text="Eliminar Seleccionada", command=eliminar_carrera, style="Danger.TButton").pack(fill=tk.X, pady=(3, 5), padx=5)

menubar = Menu(root); root.config(menu=menubar)
m_archivo = Menu(menubar, tearoff=0); menubar.add_cascade(label="Archivo", menu=m_archivo)
m_archivo.add_command(label="💾 Guardar Proyecto", command=accion_guardar_proyecto)
m_archivo.add_command(label="📂 Cargar Proyecto", command=accion_cargar_proyecto)
m_archivo.add_separator(); m_archivo.add_command(label="Salir", command=root.quit)

m_db = Menu(menubar, tearoff=0); menubar.add_cascade(label="Base de Datos", menu=m_db)
m_db.add_command(label="Importar Excel PROGRAMA", command=accion_importar_programa); m_db.add_command(label="Importar Excel RESULTADOS", command=accion_importar_resultados); m_db.add_separator(); m_db.add_command(label="⚠️ Resetear DB", command=accion_reset_db)

w, h = 1250, 850; root.geometry(f"{w}x{h}"); root.minsize(1100, 700); root.mainloop()